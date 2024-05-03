import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
import uuid
import os
class Format_2:
    def __init__(self, dfname, all_subjects,course , semester, shift, section, batch, passout_year,faculty_name,month):
        roman_numerals = {'1': 'I', '2': 'II', '3': 'III', '4': 'IV', '5': 'V','6': 'VI', '7': 'VII', '8': 'VIII', '9': 'IX', '10': 'X'}
        self.course = course
        self.semester = semester
        self.semester = roman_numerals[str(semester)]
        self.shift = shift
        self.shift_char = 'M' if shift == 'Morning' else 'E'
        self.section = section
        self.batch = batch
        self.passout_year = passout_year
        self.dfname = dfname
        self.all_subjects = all_subjects
        self.file_name = str(uuid.uuid4())
        self.faculty_name = faculty_name
        self.month=month
        
        self.df = None
        self.excel_df = None
        self.filtered_df = pd.DataFrame(columns=['Faculty Name', 'Subject', 'Appeared', 'Passed',
                                        'Pass %', '90% & Above', '75% to less than 90%', '60% to less than 75%', '50% to less than 60%', '40% to less than50%', 'Below 40%', 'Highest Marks'])
        self.excel_file_path = None
        self.word_file_path = None
        self.subject_name_mapping = {'S.No': 'S.No',
                                     'Enrollment Number': 'Enrollment Number',
                                     'Name': 'Name',
                                     'Sec': 'Sec',
                                     'Total Marks': 'Total Marks',
                                     'CGPA%': 'CGPA%',
                                     'Reappear': 'Reappear',
                                     'Absent': 'Absent', }
        for key, value in self.all_subjects.items():
            self.subject_name_mapping.update(
                {key: key+" "+value+" (Internal)"})
            self.subject_name_mapping.update(
                {key + ".1": key+" "+value + " (External)"})
            self.subject_name_mapping.update(
                {key + ".2": key+" "+value + " (Total)"})


    def read_data(self, subjects, sec):

        col_names = ['S.No', 'Enrollment Number', 'Name', 'Sec',]
        for i in range(len(subjects)):
            col_names += [subjects[i], subjects[i]+".1", subjects[i]+".2"]
        col_names += [
            'Total Marks', 'CGPA%',  'Reappear', 'Absent',]
        self.df = self.dfname
        self.df.columns = col_names
        self.df.columns = [self.subject_name_mapping.get(
            col) for col in self.df.columns]
        
        columns = ['Sec']
        for subject in subjects:
            columns.append(self.subject_name_mapping[subject + '.2'])

   
        self.excel_df = self.df[columns]
        self.excel_df = self.excel_df[self.excel_df['Sec'] == sec]
        self.excel_df = self.excel_df.drop(columns=['Sec'])

   

    def read_from_filtered_excel(self, course_name, subject_teacher_mapping):
        self.filtered_df = pd.DataFrame()  
        subject_teacher_mapping = dict(sorted(subject_teacher_mapping.items()))
        subject_code = list(subject_teacher_mapping.keys())
        teacher_list = list(subject_teacher_mapping.values())
        column_names = list()
        sum_a = sum_b = sum_students_appeared = 0
        for i in range(len(subject_code)):
            column_names.append(
                subject_code[i]+" "+self.all_subjects[subject_code[i]]+" (Total)")
        for i, column_name in enumerate(column_names):
            self.filtered_df.at[i, 'S.No'] = f'{i+1:.0f}'
            self.filtered_df.at[i, 'Faculty Name'] = teacher_list[i]
            self.filtered_df.at[i, 'Subject'] = str(
                column_name[7:-8])+f' ({column_name[:6]})'

            total_students = self.excel_df[self.excel_df[column_name]
                                           != "0"].shape[0]
            self.filtered_df.at[i, 'Appeared'] = f'{(total_students):.0f}'
            marks_list = self.excel_df[column_name].tolist()
            marks_list = [int(i) for i in marks_list]
            passed_students = len([i for i in marks_list if i >= 40])
            self.filtered_df.at[i, 'Passed'] = f'{(passed_students):.0f}'
            self.filtered_df.at[i,
                                'Pass %'] = f"{(passed_students / total_students) * 100:.2f}"
            countA1 = len([i for i in marks_list if i >= 90])
            countA2 = len([i for i in marks_list if i >= 75 and i < 90])
            countA3 = len([i for i in marks_list if i >= 60 and i < 75])
            sum_a += countA1+countA2+countA3
            sum_students_appeared += total_students
            countB1 = len([i for i in marks_list if i >= 50 and i < 60])
            countB2 = len([i for i in marks_list if i >= 40 and i < 50])
            countB3 = len([i for i in marks_list if  i>0 and i < 40])
            sum_b += countB1+countB2+countB3
            self.filtered_df.at[i, '90% & Above'] = str(
                f"{countA1}\n({countA1 / total_students * 100:.2f})")
            self.filtered_df.at[i,
                                '75% to less than 90%'] = f"{countA2}\n({countA2 / total_students * 100:.2f})"
            self.filtered_df.at[i,
                                '60% to less than 75%'] = f"{countA3}\n({countA3 / total_students * 100:.2f})"
            self.filtered_df.at[i,
                                '50% to less than 60%'] = f"{countB1}\n({countB1 / total_students * 100:.2f})"
            self.filtered_df.at[i,
                                '40% to less than50%'] = f"{countB2}\n({countB2 / total_students * 100:.2f})"
            self.filtered_df.at[i,
                                'Below 40%'] = f"{countB3}\n({countB3 / total_students * 100:.2f})"
            self.filtered_df.at[i,
                                'Highest Marks'] = f"{max(marks_list):.0f}"
            
        second_last_row = pd.Series({
            "S.No": "",
            "Faculty Name": "",
            "Subject": "Total Subjects",
            "Appeared": "Total Students Appeared",
            "Passed": '',
            "Pass %": '',
            '90% & Above': 'No. & % of Students above 60%',
            "75% to less than 90%": "",
            "60% to less than 75%": "",
            "50% to less than 60%": 'No. & % of Students below 60%',
            "40% to less than50%": "",
            "Below 40%": "",

            "Highest Marks": "",
        })
        last_row = pd.Series({
            "S.No": "",
            "Faculty Name": '',
            "Subject": len(subject_code),
            "Appeared": sum_students_appeared,
            "Passed": "",
            "Pass %": "",
            '90% & Above': f"{sum_a:.0f} ({sum_a / (sum_a+sum_b)   * 100:.2f})",
            "75% to less than 90%": "",
            "60% to less than 75%": "",
            "50% to less than 60%": f"{sum_b:.0f} ({sum_b/(sum_a+sum_b) * 100:.2f})",
            "40% to less than50%": "",
            "Below 40%": "",

            "Highest Marks": "",
        })


        self.filtered_df = self.filtered_df._append(
            second_last_row, ignore_index=True)
        self.filtered_df = self.filtered_df._append(
            last_row, ignore_index=True)
        

    def write_to_doc(self):
        self.word_file_path = os.path.join(os.path.dirname(__file__), "buffer_files", f"{self.file_name}.docx")
        excel_data_df = self.filtered_df
        doc = Document()
        for section in doc.sections:

            section.left_margin = section.right_margin = Inches(0.4)

            section.top_margin = section.bottom_margin = Inches(0)
        bold_text = ['Total Subjects','Total Students Appeared',

                     "No. & % of Students above 60%",

                     "No. & % of Students below 60%"]

        header_lines = [
            "",
            'MAHARAJA SURAJMAL INSTITUTE',
            'DEPARTMENT OF ____________',
            'Class-wise Result Analysis',
            '                   (Based on Internal & External marks)                                     Date: …………',
            f"Programme: {self.course}   Class:{self.semester} Semester - Sec-{self.section} 	        Shift: {self.shift} 	 Batch: {self.batch}	",
            f"Max. Marks: 100							            Session: {self.month} YYYY",
        ]
        
        
        



        for line in header_lines:
            paragraph = doc.add_heading(line)
            if line == '                   (Based on Internal & External marks)                                     Date: …………':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif line == 'Result Analysis (MMM-MMM YYYY)':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif line == f"Faculty Name: - Dr. ABC                        Shift-{self.shift}                                Max Marks: 100 ":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(
                    20) if paragraph.text == 'MAHARAJA SURAJMAL INSTITUTE' else Pt(14)
                font.bold = True
                if paragraph.text == 'Class-wise Result Analysis':
                    font.underline = True
                font.color.rgb = RGBColor(0, 0, 0) 

        doc.add_table(
            excel_data_df.shape[0]+1, excel_data_df.shape[1], style='Table Grid')
        table = doc.tables[0]
        
        for j in range(excel_data_df.shape[-1]):
            cell = table.cell(0, j)
            cell.text = excel_data_df.columns[j]
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)


        for i in range(excel_data_df.shape[0]):
            for j in range(excel_data_df.shape[-1]):
                cell = table.cell(i+1, j)
                cell.text = str(excel_data_df.values[i, j])
                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                if cell.text in bold_text:
                    cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        def merge_cells(table, start_row, start_col, end_row, end_col):
            for row in range(start_row, end_row + 1):
                for col in range(start_col, end_col + 1):
                    cell = table.cell(row, col)
                    if row == start_row and col == start_col:
                        continue  
                    cell.text = cell.text.strip()

    
            table.cell(start_row, start_col).merge(
                table.cell(end_row, end_col))



        merge_cells(table, start_row=len(self.filtered_df)-1, start_col=3,
                    end_row=len(self.filtered_df)-1, end_col=5)  
        merge_cells(table, start_row=len(self.filtered_df)-1,
                    start_col=6, end_row=len(self.filtered_df)-1, end_col=8)
        merge_cells(table, start_row=len(self.filtered_df)-1,
                    start_col=9, end_row=len(self.filtered_df)-1, end_col=11)
        merge_cells(table, start_row=len(self.filtered_df), start_col=3, end_row=len(
            self.filtered_df), end_col=5)  
        merge_cells(table, start_row=len(self.filtered_df),
                    start_col=6, end_row=len(self.filtered_df), end_col=8)
        merge_cells(table, start_row=len(self.filtered_df),
                    start_col=9, end_row=len(self.filtered_df), end_col=11)


        footer_lines = ["",
                        f"         Class Coordinator	                  Result Analysis Committee		  HOD, {self.course} ({self.shift_char})",
                        f"          Dr.{self.faculty_name}	                                           Dr.	                                      Dr. ",]
        for line in footer_lines:
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(14)
                font.bold = True

                font.color.rgb = RGBColor(0, 0, 0)  

        doc.save(self.word_file_path)
        return f"{self.file_name}.docx"