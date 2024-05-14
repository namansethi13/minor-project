import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION, WD_SECTION, WD_ORIENT
from docx.shared import Pt, RGBColor
import uuid
import os


class f11:
    def __init__(self, file_data, all_subjects, faculty_name, shift, semester,passing,practical_subjects):
        self.file_data = file_data
        self.df = pd.DataFrame()
        self.faculty_name = faculty_name
        self.shift = shift
        self.all_subjects = all_subjects
        print(all_subjects)
        self.file_name = str(uuid.uuid4())
        self.semester = int(semester)
        self.practical_subjects = practical_subjects#["020171","020173","020175"]
        
        if semester % 2 == 0:
            self.semester_month = "Jan-Jun"
        else:
            self.semester_month = "Jul-Dec"

    def write_to_doc(self):
        self.word_file_path = os.path.join(os.path.dirname(
            __file__), "buffer_files", f"{self.file_name}.docx")
        sub_count = row_count = 0
        doc = Document()

        def change_orientation():
            current_section = doc.sections[-1]
            new_width, new_height = current_section.page_height, current_section.page_width
            new_section = doc.add_section(WD_SECTION.NEW_PAGE)
            current_section.orientation = WD_ORIENT.LANDSCAPE
            current_section.page_width = new_width
            current_section.page_height = new_height
            return current_section
        change_orientation()
        for section in doc.sections:

            section.left_margin = section.right_margin = Inches(0.2)

            section.top_margin = section.bottom_margin = Inches(0)

        header_lines = [
            "",
            'Maharaja Surajmal Institute',
            'Department of _______',
            'Date: …………',
            f"Faculty Name: - Dr. {self.faculty_name}                        Shift-{self.shift}                                Max Marks: 100 ",
            f"Result Analysis ({self.semester_month} YYYY)"
        ]

        for line in header_lines:
            paragraph = doc.add_heading(line)
            if line == 'Date: …………':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif line == 'Result Analysis (MMM-MMM YYYY)':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif line == f"Faculty Name: - Dr. {self.faculty_name}                        Shift-{self.shift}                                Max Marks: 100 ":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(
                    20) if paragraph.text == 'Maharaja Surajmal Institute' else Pt(14)
                font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)
        
        for dfname, data in self.file_data.items():
            for section, subjects in data["section-subject"].items():
                row_count += len(subjects)
        subjects = []
        for dfname, data in self.file_data.items():
            for section, subject in data["section-subject"].items():
                for sub in subject:
                    subjects.append(sub)    
        table = doc.add_table(rows=2+row_count, cols=23)
        table.style = 'TableGrid'
        first_row = ['S.No', 'Paper Code', 'Subjects Taught', 'Students Appeared', '', 'Passed', '', 'Pass%', '',
                     '>=90%', '', "89.99-75%", '', "74.99-60%", '', "59.99-50%", '', "49.99-40%", '', "<40%", '', "Highest Marks"]
        second_row = ["", "", ""] + ['Int', 'Ext']*10
        for i in range(len(first_row)):
            table.cell(0, i).text = first_row[i]

        for i in range(3):
            table.cell(0, i).merge(table.cell(1, i))
            table.cell(0, i).text = first_row[i]
        for i in range(3, len(first_row), 2):
            table.cell(0, i).merge(table.cell(0, i+1))
            table.cell(0, i).text = first_row[i]
        for i in range(3, len(second_row)):
            table.cell(1, i).text = second_row[i]
        for dfname, data in self.file_data.items():
            all_columns = []
        
            print(data["all_columns"])
            for i in range(len(data["all_columns"])):
                all_columns.append(data["all_columns"][i])
                all_columns.append(data["all_columns"][i]+"_ext")
            result_df = data["result_df"]
            for section, subjects in data["section-subject"].items():
                for i in range(len(subjects)):
                    self.df = result_df
                    self.df = self.df.iloc[:, :-4]
                    need = []
                    for j in range(4, len(self.df.columns)):
                        if j % 3 != 0:
                            need.append(j)

                    self.df = self.df.iloc[:, [0, 1, 2, 3]+need]
                    self.df.columns = ['S.No', 'Name',
                                    'Enrollment No', 'Section']+all_columns

                    self.df = self.df[self.df['Section'] == section]
                    
                    self.df = self.df.iloc[:, :4].join(
                        self.df[subjects[i]]).join(self.df[subjects[i]+"_ext"])

                    non_empty_values1 = self.df[subjects[i]].dropna()
                    non_empty_values2 = self.df[subjects[i]+"_ext"].dropna()

                    marks_list1 = non_empty_values1.to_list()
                    marks_list1 = [int(i) for i in marks_list1]
                    
                    marks_list2 = non_empty_values2.to_list()
                    marks_list2 = [int(i) for i in marks_list2]
                    

                    self.df = pd.to_numeric(self.df[subjects[i]], errors='coerce')
                        
                    table.cell(sub_count+2, 0).text = str(sub_count+1)
                    table.cell(
                        sub_count+2, 1).text = data["course"]+str(subjects[i])[-3:] + section
                    table.cell(
                            sub_count+2, 2).text = self.all_subjects[subjects[i]]
                    if subjects[i] not in self.practical_subjects:
                        subjects[i] = int(subjects[i])
                        table.cell(
                            sub_count+2, 3).text = str(len([i for i in marks_list1 if i > 0]))
                        table.cell(
                            sub_count+2, 4).text = str(len([i for i in marks_list2 if i > 0]))

                        table.cell(
                            sub_count+2, 5).text = str(len([i for i in marks_list1 if i >= 10]))
                        table.cell(
                            sub_count+2, 6).text = str(len([i for i in marks_list2 if i >= 30]))
                        table.cell(
                            sub_count+2, 7).text = f"{len([i for i in marks_list1 if i>=10])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 8).text = f"{len([i for i in marks_list2 if i>=30])/len([i for i in marks_list2 if i > 0])*100:.2f}%"

                        table.cell(
                            sub_count+2, 9).text = f"{len([i for i in marks_list1 if i>= 22.5])}\n({len([i for i in marks_list1 if i>= 22.5])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 10).text = f"{len([i for i in marks_list2 if i>= 67.5])}\n({len([i for i in marks_list2 if i>= 67.5])/len([i for i in marks_list2 if i > 0])*100:.2f}%"

                        table.cell(
                            sub_count+2, 11).text = f"{len([ i for i in marks_list1 if i>=18.75 and i<22.5])}\n({len([i for i in marks_list1 if i>=18.75 and i<22.5 ])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 12).text = f"{len([ i for i in marks_list2 if i>=56.25 and i<67.5])}\n({len([i for i in marks_list2 if i>=56.25 and i<67.5])/len([i for i in marks_list2 if i > 0])*100:.2f}%"

                        table.cell(
                            sub_count+2, 13).text = f"{len([i for i in marks_list1 if i >= 15 and i <18.75])}\n({len([i for i in marks_list1 if i >= 60 and i <18.75])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 14).text = f"{len([i for i in marks_list2 if i >= 45 and i <56.25])}\n({len([i for i in marks_list2 if i >= 45 and i <56.25])/len([i for i in marks_list2 if i > 0])*100:.2f}%"

                        table.cell(
                            sub_count+2, 15).text = f"{len([i for i in marks_list1 if i >= 12.5 and i <15])}\n({len([i for i in marks_list1 if i >= 12.5 and i <15])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 16).text = f"{len([i for i in marks_list2 if i >= 37.5 and i <45])}\n({len([i for i in marks_list2 if i >= 37.5 and i <45])/len([i for i in marks_list2 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 17).text = f"{len([i for i in marks_list1 if i >= 10 and i <12.5])}\n({len([i for i in marks_list1 if i >= 10 and i <12.5])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 18).text = f"{len([i for i in marks_list2 if i >= 30 and i <37.5])}\n({len([i for i in marks_list2 if i >= 30 and i <37.5])/len([i for i in marks_list2 if i > 0])*100:.2f}%"

                        table.cell(
                            sub_count+2, 19).text = f"{len([i for i in marks_list1 if i > 0 and i < 10])}\n({len([i for i in marks_list1 if i > 0 and i < 10])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 20).text = f"{len([i for i in marks_list2 if i > 0 and i < 30])}\n({len([i for i in marks_list2 if i > 0 and i < 30])/len([i for i in marks_list2 if i > 0])*100:.2f}%"
                        table.cell(sub_count+2, 21).text = f"{max(marks_list1)}"
                        table.cell(sub_count+2, 22).text = f"{max(marks_list2)}"
                        
                    elif subjects[i] in self.practical_subjects:
                        subjects[i] = int(subjects[i])
                        table.cell(
                            sub_count+2, 3).text = str(len([i for i in marks_list1 if i > 0]))
                        table.cell(
                            sub_count+2, 4).text = str(len([i for i in marks_list2 if i > 0]))

                        table.cell(
                            sub_count+2, 5).text = str(len([i for i in marks_list1 if i >= 16]))
                        table.cell(
                            sub_count+2, 6).text = str(len([i for i in marks_list2 if i >= 24]))
                        table.cell(
                            sub_count+2, 7).text = f"{len([i for i in marks_list1 if i>=16])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 8).text = f"{len([i for i in marks_list2 if i>=24])/len([i for i in marks_list2 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 9).text = f"{len([i for i in marks_list1 if i>= 36])}\n({len([i for i in marks_list1 if i>= 36])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 10).text = f"{len([i for i in marks_list2 if i>= 54])}\n({len([i for i in marks_list2 if i>= 54])/len([i for i in marks_list2 if i > 0])*100:.2f}%"

                        table.cell(
                            sub_count+2, 11).text = f"{len([ i for i in marks_list1 if i>=30 and i<36])}\n({len([i for i in marks_list1 if i>=30 and i<36 ])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 12).text = f"{len([ i for i in marks_list2 if i>=45 and i<54])}\n({len([i for i in marks_list2 if i>=45 and i<54])/len([i for i in marks_list2 if i > 0])*100:.2f}%"

                        table.cell(
                            sub_count+2, 13).text = f"{len([i for i in marks_list1 if i >= 24 and i <30])}\n({len([i for i in marks_list1 if i >= 24 and i <30])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 14).text = f"{len([i for i in marks_list2 if i >= 36 and i <45])}\n({len([i for i in marks_list2 if i >= 36 and i <45])/len([i for i in marks_list2 if i > 0])*100:.2f}%"
                        
                        table.cell(
                            sub_count+2, 15).text = f"{len([i for i in marks_list1 if i >= 20 and i <24])}\n({len([i for i in marks_list1 if i >= 20 and i <24])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 16).text = f"{len([i for i in marks_list2 if i >= 30 and i <36])}\n({len([i for i in marks_list2 if i >= 30 and i <36])/len([i for i in marks_list2 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 17).text = f"{len([i for i in marks_list1 if i >= 16 and i <20])}\n({len([i for i in marks_list1 if i >= 16 and i <20])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 18).text = f"{len([i for i in marks_list2 if i >= 24 and i <30])}\n({len([i for i in marks_list2 if i >= 24 and i <30])/len([i for i in marks_list2 if i > 0])*100:.2f}%"
                        
                        table.cell(
                            sub_count+2, 19).text = f"{len([i for i in marks_list1 if i > 0 and i < 16])}\n({len([i for i in marks_list1 if i > 0 and i < 16])/len([i for i in marks_list1 if i > 0])*100:.2f}%"
                        table.cell(
                            sub_count+2, 20).text = f"{len([i for i in marks_list2 if i > 0 and i < 24])}\n({len([i for i in marks_list2 if i > 0 and i < 24])/len([i for i in marks_list2 if i > 0])*100:.2f}%"
                        table.cell(sub_count+2, 21).text = f"{max(marks_list1)}"
                        table.cell(sub_count+2, 22).text = f"{max(marks_list2)}"
                    sub_count += 1
                
                    
        for i in range(0, row_count+2):
            for j in range(23):
                table.cell(
                    i, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in table.cell(i, j).paragraphs[0].runs:
                    font = run.font
                    font.name = 'Times New Roman'
                    font.size = Pt(10)
                    font.bold = True
                    font.color.rgb = RGBColor(0, 0, 0)

        footer_lines = [
            '',
            '“I do hereby solemnly affirm and declare that the facts stated in the above result are true to the best of my knowledge and belief”',
            f"""Dr.{self.faculty_name}    		                 (Dr.ABC)       				(Mr. ABC)	
Assistant Professor 		Convenor-Result Analysis Committee         	HOD-____"""
        ]
        for line in footer_lines:
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(14)
                if line != '“I do hereby solemnly affirm and declare that the facts stated in the above result are true to the best of my knowledge and belief”':
                    font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)
        doc.save(self.word_file_path)
        return f"{self.file_name}.docx"
