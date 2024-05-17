import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Pt, RGBColor
import uuid
import os
import uuid


class f13:
    def __init__(self, request_data, shift, year, faculty_name, semester, section, is_practical):
        self.is_practical = is_practical
        course = request_data["course"]
        self.course = course
        self.shift = shift
        self.semester = semester
        self.section = section
        year = request_data["year"]
        subject = request_data["subject"]
        self.df = request_data["df"]
        last_three_columns = self.df.columns[-3:]
        self.df.dropna(inplace=True)
        self.df.sort_values(
            by=last_three_columns[-3], ascending=False, inplace=True)
        self.df.insert(0, 'S.No.', range(1, 1 + len(self.df)))
        print('df:', self.df.columns)
        self.subject_name = " ".join(self.df.columns[-1].split(" ")[:-1])
        print('subject_name:', self.subject_name)
        self.file_name = str(uuid.uuid4())
        self.faculty_name = faculty_name
        self.shift = shift
        self.result_year = year
        roman_map = {
            1: 'I',
            2: 'II',
            3: 'III',
            4: 'IV',
            5: 'V',
            6: 'VI',
            7: 'VII',
            8: 'VIII',
            9: 'IX',
            10: 'X'
        }
        self.semester_roman = roman_map[int(semester)]
        str_enrollment_series = self.df['Enrollment Number'].tolist()
        print(str_enrollment_series)

        for i in range(len(str_enrollment_series)):
            str_enrollment_series[i] = f"{int(str_enrollment_series[i]):110d}"

    def write_to_doc(self):
        self.word_file_path = os.path.join(os.path.dirname(
            __file__), "buffer_files", f"{self.file_name}.docx")
        doc = Document()
        header_lines = [
            "",
            'MAHARAJA SURAJMAL INSTITUTE',
            f"DEPARTMENT OF _________________({self.shift})",
            f"                      {self.shift} Shift       MMM-MMM YYYY",
            f"Class:- {self.course} {self.semester_roman} Sec {self.section} Batch [yyyy-yyyy]",
            f"(Note: Student list is sorted in decreasing order based on internal marks)",
            f"Dr. {self.faculty_name}",
            f" "
        ]
        footer_lines = footer_lines = [
            '“I do hereby solemnly affirm and declare that the facts stated in the above result are true to the best of my knowledge and belief”',
            f"""Dr.{self.faculty_name}    		                 (Dr.ABC)       				(Mr. ABC)	
Assistant Professor 		Convenor-Result Analysis Committee         	HOD-____"""
        ]
        for line in header_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(
                    20) if paragraph.text == 'Maharaja Surajmal Institute' else Pt(14)
                font.size = Pt(12) if "(Note:" in paragraph.text else Pt(14)
                font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)
        doc.add_table(rows=self.df.shape[0]+2, cols=self.df.shape[1])
        table = doc.tables[0]
        table.style = 'TableGrid'
        table.cell(0, 0).merge(table.cell(1, 0)).text = "S.No."
        table.cell(0, 1).merge(table.cell(1, 1)).text = "Enrollment No."
        table.cell(0, 2).merge(table.cell(1, 2)).text = "Name"
        table.cell(0, 3).merge(table.cell(0, 5)).text = self.subject_name
        table.cell(1, 3).text = "Int"
        table.cell(1, 4).text = "Ext"
        table.cell(1, 5).text = "Total"
        sno = list(self.df.iloc[:, 0])
        enroll_no = list(self.df.iloc[:, 1])
        name = list(self.df.iloc[:, 2])
        int_marks = list(self.df.iloc[:, -3])
        ext_marks = list(self.df.iloc[:, -2])
        total_marks = list(self.df.iloc[:, -1])
        for i in range(self.df.shape[0]):
            table.cell(i+2, 0).text = str(sno[i])
            table.cell(i+2, 1).text = f"{int(enroll_no[i]):110d}"
            table.cell(i+2, 2).text = str(name[i])
            table.cell(i+2, 3).text = str(int_marks[i])
            table.cell(i+2, 4).text = str(ext_marks[i])
            table.cell(i+2, 5).text = str(total_marks[i])
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(0)
                    paragraph_format.space_after = Pt(0)
                    paragraph_format.line_spacing = 1
                    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = Pt(9)
                        font.color.rgb = RGBColor(0, 0, 0)

        doc.add_paragraph()

        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(
                    20) if paragraph.text == 'MAHARAJA SURAJMAL INSTITUTE' else Pt(14)
                font.bold = True
                if paragraph.text == 'Class-wise Result Analysis':
                    font.underline = True
                font.color.rgb = RGBColor(0, 0, 0)
        doc.add_page_break()
        doc.add_table(rows=3, cols=18)
        table2 = doc.tables[1]
        table2.style = 'TableGrid'
        first_row = ["Student Appeared", "Student Passed", "Pass %", ">=90%",
                     "89.99-75%", "74.99-60%", "59.99-50%", "49.99-40%", "<40%"]
        second_row = ["Int", "Ext"]*(len(first_row))

        for i in range(len(int_marks)):
            int_marks[i] = 0 if int_marks[i] == 'AB' else int(int_marks[i])
            ext_marks[i] = 0 if ext_marks[i] == 'AB' else int(ext_marks[i])
        for i in range(len(first_row)):
            table2.cell(0, i*2).merge(table2.cell(0, i*2+1)
                                      ).text = first_row[i]
            table2.cell(1, i*2).text = second_row[i*2]
            table2.cell(1, i*2+1).text = second_row[i*2+1]
        table2.cell(2, 0).text = f"{len([i for i in int_marks if i != 0])}"
        table2.cell(2, 1).text = f"{len([i for i in ext_marks if i != 0])}"
        if self.is_practical:
            table2.cell(
                2, 2).text = f"{len([i for i in int_marks if i > 16])}"
            table2.cell(
                2, 3).text = f"{len([i for i in ext_marks if i > 24])}"
            table2.cell(
                2, 4).text = f"{len([i for i in int_marks if i > 16])/len([i for i in int_marks if i != 0])*100:.2f}%"
            table2.cell(
                2, 5).text = f"{len([i for i in ext_marks if i > 24])/len([i for i in ext_marks if i != 0])*100:.2f}%"
            table2.cell(
                2, 6).text = f"{len([i for i in int_marks if i >= 36])} ({len([i for i in int_marks if i >= 36])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 7).text = f"{len([i for i in ext_marks if i >= 54])} ({len([i for i in ext_marks if i >= 54])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 8).text = f"{len([i for i in int_marks if 30 <= i < 36])} ({len([i for i in int_marks if 30 <= i < 36])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 9).text = f"{len([i for i in ext_marks if 45 <= i < 54])} ({len([i for i in ext_marks if 45 <= i < 54])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 10).text = f"{len([i for i in int_marks if 24 <= i < 30])} ({len([i for i in int_marks if 24 <= i < 30])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 11).text = f"{len([i for i in ext_marks if 36 <= i < 45])} ({len([i for i in ext_marks if 36 <= i < 45])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 12).text = f"{len([i for i in int_marks if 20 <= i < 24])} ({len([i for i in int_marks if 20 <= i < 24])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 13).text = f"{len([i for i in ext_marks if 30 <= i < 36])} ({len([i for i in ext_marks if 30 <= i < 36])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 14).text = f"{len([i for i in int_marks if 16 <= i < 20])} ({len([i for i in int_marks if 16 <= i < 20])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 15).text = f"{len([i for i in ext_marks if 24 <= i < 30])} ({len([i for i in ext_marks if 24 <= i < 30])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 16).text = f"{len([i for i in int_marks if i < 16])} ({len([i for i in int_marks if i < 16])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 17).text = f"{len([i for i in ext_marks if i < 24])} ({len([i for i in ext_marks if i < 24])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
        else:
            table2.cell(
                2, 2).text = f"{len([i for i in int_marks if i > 10])}"

            table2.cell(
                2, 3).text = f"{len([i for i in ext_marks if i > 30])}"
            table2.cell(
                2, 4).text = f"{len([i for i in int_marks if i > 10])/len([i for i in int_marks if i != 0])*100:.2f}%"
            table2.cell(
                2, 5).text = f"{len([i for i in ext_marks if i > 30])/len([i for i in ext_marks if i != 0])*100:.2f}%"
            table2.cell(
                2, 6).text = f"{len([i for i in int_marks if i >= 22.5])} ({len([i for i in int_marks if i >= 22.5])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 7).text = f"{len([i for i in ext_marks if i >= 67.5])} ({len([i for i in ext_marks if i >= 67.5])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 8).text = f"{len([i for i in int_marks if 18.75 <= i < 22.5])} ({len([i for i in int_marks if 18.75 <= i < 22.5])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 9).text = f"{len([i for i in ext_marks if 56.25 <= i < 67.5])} ({len([i for i in ext_marks if 56.25 <= i < 67.5])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 10).text = f"{len([i for i in int_marks if 15 <= i < 18.75])} ({len([i for i in int_marks if 15 <= i < 18.75])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 11).text = f"{len([i for i in ext_marks if 45 <= i < 56.25])} ({len([i for i in ext_marks if 45 <= i < 56.25])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 12).text = f"{len([i for i in int_marks if 12.5 <= i < 15])} ({len([i for i in int_marks if 12.5 <= i < 15])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 13).text = f"{len([i for i in ext_marks if 37.5 <= i < 45])} ({len([i for i in ext_marks if 37.5 <= i < 45])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 14).text = f"{len([i for i in int_marks if 10 <= i < 12.5])} ({len([i for i in int_marks if 10 <= i < 12.5])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 15).text = f"{len([i for i in ext_marks if 30 <= i < 37.5])} ({len([i for i in ext_marks if 30 <= i < 37.5])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 16).text = f"{len([i for i in int_marks if i < 10])} ({len([i for i in int_marks if i < 10])/len([i for i in int_marks if i != 0])*100:.2f}%)"
            table2.cell(
                2, 17).text = f"{len([i for i in ext_marks if i < 30])} ({len([i for i in ext_marks if i < 30])/len([i for i in ext_marks if i != 0])*100:.2f}%)"
        for row in table2.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(0)
                    paragraph_format.space_after = Pt(0)
                    paragraph_format.line_spacing = 1
                    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = Pt(10)
                        font.color.rgb = RGBColor(0, 0, 0)

        for line in footer_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(
                    7) if "I do hereby " in paragraph.text else Pt(9)
                font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)
        doc.save(self.word_file_path)
        return f"{self.file_name}"
