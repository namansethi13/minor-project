import os
import uuid
import math
import docx
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Pt, RGBColor
roman_numerals = {'1': 'I', '2': 'II', '3': 'III', '4': 'IV',
                  '5': 'V', '6': 'VI', '7': 'VII', '8': 'VIII', '9': 'IX', '10': 'X'}


class Format6:
    def __init__(self, df_data, all_subjects, shift, course, month, faculty_name, admitted_years, all_semesters):
        self.df_data = df_data
        self.result_years = []
        for dfname, data in self.df_data.items():
            data["subjects"] = ["SNo", "Enrollment No.",
                                "Name", "Section"] + data["subjects"]
        self.df = pd.DataFrame()
        self.filtered_df = pd.DataFrame()
        self.all_subjects = all_subjects
        self.file_name = str(uuid.uuid4())
        self.shift = shift
        self.course = course
        self.month = month

        self.faculty_name = faculty_name
        self.admitted_years = admitted_years
        for i, a in enumerate(self.admitted_years):
            result_year = int(a) + math.ceil(int(all_semesters[i])/2)

            if not int(all_semesters[i]) % 2 == 0:
                result_year = result_year - 1

            self.result_years.append(result_year)

    def write_to_doc(self):
        self.word_file_path = os.path.join(os.path.dirname(
            __file__), "buffer_files", f"{self.file_name}.docx")
        table_count = 0
        doc = Document()
        for section in doc.sections:

            section.left_margin = section.right_margin = Inches(0.4)

            section.top_margin = section.bottom_margin = Inches(0)
        header_lines = [
            "",
            'MAHARAJA SURAJMAL INSTITUTE',
            f"DEPARTMENT OF _________________({self.shift})",
            f"                      Faculty Name: Dr. {self.faculty_name}       {self.month} {self.result_years[0]}              Date: …………",
        ]
        footer_lines = [f"Faculty                         Result Analysis Committee	             HOD,{self.course} ({self.shift})",
                        f"{self.faculty_name}"]
        for line in header_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        iterator = 0
        for dfname, data in self.df_data.items():

            column_names = data["subjects"]
            needed_subjects = data["needed_subjects"]
            sections = data["sections"]
            course = data["course"]
            shift = data["shift"]
            semester = data["semester"]
            resut_df = data["result_df"]
            result_year = self.result_years[iterator]

            iterator += 1

            for i in range(len(data["sections"])):

                self.df = resut_df.copy()

                self.df = self.df.iloc[:, :-4]
                self.df = self.df.iloc[:, [
                    0, 1, 2, 3]+[i for i in range(6, len(self.df.columns), 3)]]
                self.df.columns = column_names

                self.df = self.df[self.df['Section'] == sections[i]]

                self.df = self.df[self.df[needed_subjects[i]] != 0]
                self.df = self.df.iloc[:, :4].join(self.df[needed_subjects[i]])
                self.topdf = self.df.sort_values(
                    by=needed_subjects[i], ascending=False).head(10)
                self.bottomdf = self.df.sort_values(
                    by=needed_subjects[i], ascending=True).head(10)
                self.bottomdf.to_csv('bottom.csv')
                doc.add_paragraph()
                table_count += 1
                if table_count != 1:
                    if table_count % 2 == 1:
                        doc.add_page_break()

                doc.add_table(rows=13, cols=7)

                doc.page_break_before = True
                table = doc.tables[-1]
                table.style = "Table Grid"
                table.autofit = True
                table.cell(
                    0, 1).text = f'Subject Name: {self.all_subjects[needed_subjects[i]]}    Class: {course} {roman_numerals[str(semester)]} Semester (Section {sections[i]})     Shift {shift}'

                table.cell(0, 1).merge(table.cell(0, 6))

                table.cell(
                    1, 1).text = f"Top 10 Students({self.month} {result_year})"

                table.cell(
                    1, 4).text = f"Bottom 10 Students ({self.month} {result_year})"

                table.cell(1, 1).merge(table.cell(1, 3))
                table.cell(1, 4).merge(table.cell(1, 6))

                table.cell(2, 0).text = "S.no"
                table.cell(2, 1).text = "Enrol No."
                table.cell(2, 2).text = "Name"
                table.cell(2, 3).text = "Marks"
                table.cell(2, 4).text = "Enrol No."
                table.cell(2, 5).text = "Name"
                table.cell(2, 6).text = "Marks"
                self.df.to_csv('temp.csv')
                for j in range(10):
                    table.cell(3+j, 0).text = str(j+1)
                    table.cell(
                        3+j, 1).text = f"{int(self.topdf.iloc[j, 1]):011d}"
                    table.cell(
                        3+j, 2).text = str(self.topdf.iloc[j, 2]).casefold().title()
                    table.cell(
                        3+j, 3).text = str(int(self.topdf.iloc[j, 4]))
                    table.cell(
                        3+j, 4).text = f"{int(self.bottomdf.iloc[j,1]):011d}"
                    table.cell(
                        3+j, 5).text = str(self.bottomdf.iloc[j, 2]).casefold().title()
                    table.cell(
                        3+j, 6).text = str(int(self.bottomdf.iloc[j, 4]))

                for row in table.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.space_before = Pt(0)
                            paragraph_format.space_after = Pt(0)
                            paragraph_format.line_spacing = 1
                            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            for run in paragraph.runs:
                                font = run.font
                                font.name = 'Times New Roman'
                                font.size = Pt(11)
                                font.bold = True
                                font.color.rgb = RGBColor(0, 0, 0)

            doc.add_paragraph()

        for line in footer_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(
                    20) if paragraph.text == 'MAHARAJA SURAJMAL INSTITUTE' else Pt(14)
                font.bold = True
                if paragraph.text == 'Class-wise Result Analysis':
                    font.underline = True
                font.color.rgb = RGBColor(0, 0, 0)

        doc.save(self.word_file_path)
        return f"{self.file_name}.docx"
