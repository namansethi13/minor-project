import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION, WD_SECTION, WD_ORIENT
from docx.shared import Pt, RGBColor
import uuid
import math
import os

class Format7:
    def __init__(self,dataframe, file_data,faculty_name,all_subjects,admitted):
        self.faculty_name = faculty_name
        self.file_data = file_data
        self.semester = file_data["semester"]
        self.course = file_data["course"]
        self.shift_char = file_data["shift"]
        if self.semester%2==0:
            self.months = "Jan-July"
        else:
            self.months = "Aug-Dec"
        self.admitted_year = admitted
        self.result_year = int(self.admitted_year) + math.ceil(self.semester/2)
        if not self.semester%2==0:
            self.result_year = self.result_year - 1
        self.dataframe = dataframe
        self.file_data["Subjects"] = ["SNo","Enrollment No.", "Name","Section"] + self.file_data["Subjects"]+["CGPA%"]
        self.df = pd.DataFrame()
        self.filtered_df = pd.DataFrame()
        self.all_subjects = all_subjects
        self.file_name = str(uuid.uuid4())

    def write_to_doc(self):
        self.word_file_path = os.path.join(os.path.dirname(__file__), "buffer_files", f"{self.file_name}.docx")
        self.df = self.dataframe.copy()
        
        
        self.df = self.df.iloc[:,:-2]
        self.df = self.df.drop(self.df.columns[-2],axis=1)
       
        self.df.rename(columns={'Unnamed: 0':'SNo',' ':'Enrollment No.','Unnamed: 2':'Name','Unnamed: 3':'Section'},inplace=True)
        
        
        self.df = self.df.sort_values(by=['CGPA%'],ascending=False)
        self.df.to_csv("test.csv")
        self.df = self.df[self.df['Sec']==self.file_data['Section']]
        
        self.df = self.df.drop('Sec',axis=1)
        
        self.df = self.df.dropna()
        
        self.topdf = self.df.iloc[:10,:]
        
        self.bottomdf = self.df.iloc[-10:,:]
        self.bottomdf=self.bottomdf.sort_values(by=['CGPA%'],ascending=True)
        doc=Document()
        def change_orientation():
            current_section = doc.sections[-1]
            new_width, new_height = current_section.page_height, current_section.page_width
            new_section = doc.add_section(WD_SECTION.NEW_PAGE)
            new_section.orientation = WD_ORIENT.LANDSCAPE
            new_section.page_width = new_width
            new_section.page_height = new_height
            return new_section

        change_orientation()
        header_lines = [
            
            'MAHARAJA SURAJMAL INSTITUTE',
            'DEPARTMENT OF _________________',
            f"                      Faculty Name: Dr. {self.faculty_name}       {self.months} {self.result_year}              Date: ",
        ]
        footer_lines = [f"Faculty                         Result Analysis Committee	             HOD, {self.course} ({self.shift_char})",
                        f"Dr. {self.faculty_name}"]
        
        for section in doc.sections:
            section.left_margin = section.right_margin = Inches(0.4)
            section.top_margin = section.bottom_margin = Inches(0)
            section.orientation = WD_ORIENTATION.LANDSCAPE
        
        for line in header_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_table(rows=13,cols=(self.file_data["Subjects"].__len__()-5)*3+4)
        table = doc.tables[-1]
        table.style = 'Table Grid'
        table.cell(0,0).text = "SNo"
        table.cell(0,1).text = "Enrollment No."
        table.cell(0,2).text = "Name of the student"
        for i in range(len(self.file_data["Faculty Names"])):

            table.cell(0,3+i*3).text = self.file_data["Faculty Names"][i]
            table.cell(0,4+i*3).text = ""
            table.cell(0,5+i*3).text = ""
            table.cell(0,3+i*3).merge(table.cell(0,5+i*3))
        for i in range(3):
            table.cell(1,i).text = ""
        for i in range(len(self.file_data["Faculty Names"])):
            print(self.file_data["Subjects"])
            table.cell(1,3+i*3).text = self.all_subjects[self.file_data["Subjects"][4+i]]
            table.cell(1,4+i*3).text = ""
            table.cell(1,5+i*3).text = ""
            table.cell(1,3+i*3).merge(table.cell(1,5+i*3))
        table.cell(1,3+len(self.file_data["Faculty Names"])*3).text = "Total"
        table.cell(1,3+len(self.file_data["Faculty Names"])*3).merge(table.cell(0,3+len(self.file_data["Faculty Names"])*3))
        table.cell(2,3+len(self.file_data["Faculty Names"])*3).text = "%"
        for i in range(3):
            table.cell(2,i).text = ""
        for i in range(len(self.file_data["Faculty Names"])):
            if self.file_data["Subjects"][4+i] in self.file_data["Practicals"]:
                table.cell(2,3+i*3).text = "Int (40)"
                table.cell(2,4+i*3).text = "Ext (60)"
                table.cell(2,5+i*3).text = "T (100)"
            else:
                table.cell(2,3+i*3).text = "Int (25)"
                table.cell(2,4+i*3).text = "Ext (75)"
                table.cell(2,5+i*3).text = "T (100)"
      
        for i in range(10):
            table.cell(i+3,0).text = str(i+1)
            table.cell(i+3,1).text = f"{self.topdf.iloc[i,1]:.0f}"
            table.cell(i+3,2).text = str(self.topdf.iloc[i,2]).title()
            for j in range(len(self.file_data["Faculty Names"])):
                table.cell(i+3, 3+j*3).text = f'{float(self.topdf.iloc[i, 3+j*3]):.0f}'  # Convert to float
                table.cell(i+3,4+j*3).text = f"{float(self.topdf.iloc[i,4+j*3]):.0f}"if self.topdf.iloc[i,4+j*3]!="0" else "A"
                table.cell(i+3,5+j*3).text = f"{float(self.topdf.iloc[i,5+j*3]):.0f}"if self.topdf.iloc[i,4+j*3]!="0" else "A"
            table.cell(i+3,3+len(self.file_data["Faculty Names"])*3).text = f"{self.topdf.iloc[i,-1]:.2f}%"
            for row in table.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.space_before = Pt(0)
                            paragraph_format.space_after = Pt(0)
                            paragraph_format.line_spacing = 1
                            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                font = run.font
                                font.name = 'Times New Roman'
                                font.size = Pt(11)
                                font.bold = True
                                font.color.rgb = RGBColor(0, 0, 0)
        
        for line in footer_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_page_break()
        
        for line in header_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_table(rows=13,cols=(self.file_data["Subjects"].__len__()-5)*3+4)
        table = doc.tables[-1]
        table.style = 'Table Grid'
        table.cell(0,0).text = "SNo"
        table.cell(0,1).text = "Enrollment No."
        table.cell(0,2).text = "Name of the student"
        for i in range(len(self.file_data["Faculty Names"])):

            table.cell(0,3+i*3).text = self.file_data["Faculty Names"][i]
            table.cell(0,4+i*3).text = ""
            table.cell(0,5+i*3).text = ""
            table.cell(0,3+i*3).merge(table.cell(0,5+i*3))
        for i in range(3):
            table.cell(1,i).text = ""
        for i in range(len(self.file_data["Faculty Names"])):
            table.cell(1,3+i*3).text = self.all_subjects[self.file_data["Subjects"][4+i]]
            table.cell(1,4+i*3).text = ""
            table.cell(1,5+i*3).text = ""
            table.cell(1,3+i*3).merge(table.cell(1,5+i*3))
        table.cell(1,3+len(self.file_data["Faculty Names"])*3).text = "Total"
        table.cell(1,3+len(self.file_data["Faculty Names"])*3).merge(table.cell(0,3+len(self.file_data["Faculty Names"])*3))
        table.cell(2,3+len(self.file_data["Faculty Names"])*3).text = "%"
        for i in range(3):
            table.cell(2,i).text = ""
        for i in range(len(self.file_data["Faculty Names"])):
            if self.file_data["Subjects"][4+i] in self.file_data["Practicals"]:
                table.cell(2,3+i*3).text = "Int (40)"
                table.cell(2,4+i*3).text = "Ext (60)"
                table.cell(2,5+i*3).text = "T (100)"
            else:
                table.cell(2,3+i*3).text = "Int (25)"
                table.cell(2,4+i*3).text = "Ext (75)"
                table.cell(2,5+i*3).text = "T (100)"
        
        for i in range(10):
            table.cell(i+3,0).text = str(i+1)
            table.cell(i+3,1).text = f"{self.bottomdf.iloc[i,1]:.0f}"
            table.cell(i+3,2).text = str(self.bottomdf.iloc[i,2]).title()
            for j in range(len(self.file_data["Faculty Names"])):
                table.cell(i+3,3+j*3).text = f"{float(self.bottomdf.iloc[i,3+j*3]):.0f}"
                table.cell(i+3,4+j*3).text = f"{float(self.bottomdf.iloc[i,4+j*3]):.0f}"if self.bottomdf.iloc[i,4+j*3]!=0 else "A"
                table.cell(i+3,5+j*3).text = f"{float(self.bottomdf.iloc[i,5+j*3]):.0f}"if self.bottomdf.iloc[i,5+j*3]!=0 else "A"
            table.cell(i+3,3+len(self.file_data["Faculty Names"])*3).text = f"{self.bottomdf.iloc[i,-1]:.2f}%"
        for line in footer_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        for row in table.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.space_before = Pt(0)
                            paragraph_format.space_after = Pt(0)
                            paragraph_format.line_spacing = 1
                            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                font = run.font
                                font.name = 'Times New Roman'
                                font.size = Pt(9)
                                font.bold = True
                                font.color.rgb = RGBColor(0, 0, 0)
        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(
                    20) if paragraph.text == 'MAHARAJA SURAJMAL INSTITUTE' else Pt(14)
                font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)
        doc.save(self.word_file_path)
        return f"{self.file_name}.docx"

        

