from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
import uuid
import os


class f4():
    def __init__(self, data):
        self.data = data
        print("data",data)
        self.batches=[i for i in data[0]['data'].keys()]
        self.file_name = str(uuid.uuid4())
        
    def write_to_doc(self):
        word_file_path = os.path.join(os.path.dirname(
            __file__), "buffer_files", f"{self.file_name}.docx")
        doc = Document()
        for i in self.data:
            print(i)
            header_lines = [
                'Maharaja Surajmal Institute',
                'Department of _______',
                f'Programme: {i["course"]}({i["shift"][0]})',
                "Date: ___________",

            ]

            for line in header_lines:
                paragraph = doc.add_heading(line)
                if line == "Date: ___________":
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for paragraph in doc.paragraphs:
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Times New Roman'
                    if paragraph.text == 'Maharaja Surajmal Institute':
                            font.size = Pt(24)  
                    elif paragraph.text == 'Department of _______':
                            font.size = Pt(20)
                    elif paragraph.text == f'Programme: {i["course"]}({i["shift"][0]})':
                            font.size = Pt(16)
                    elif paragraph.text == "Date: ___________":
                            font.size = Pt(12)
                    font.bold = True
                    font.color.rgb = RGBColor(0, 0, 0)
            total_rows=0
            for j in self.data:
                # print("j",j)
                total_rows+=len(j["data"])+1
                # print("total_rows",total_rows)
                doc.add_table(rows=total_rows, cols=5)
                table = doc.tables[0]
                table.style = 'Table Grid'
                row1=table.rows[0]
                row1.cells[0].text = "Batch"
                row1.cells[1].text = "Semester"
                row1.cells[2].text = "No. and % of students with <60%"
                row1.cells[3].text = "No. and % of students with >=60%"
                row1.cells[4].text = "No. and % of students with >=90%"
                batch_count=0
                for batch in j["data"]:
                    # print("batch",batch)
                    batch_count+=1
                    table.rows[batch_count].cells[0].text = batch
                    table.rows[batch_count].cells[1].text = batch.items()[0][1]
                    cgpa_list= j["data"][batch][str(j["semester"])].values()
                    count_greater_60=sum(cgpa >= 60 for cgpa in cgpa_list)
                    count_less_60=sum(cgpa < 60 for cgpa in cgpa_list)
                    count_greater_90=sum(cgpa >= 90 for cgpa in cgpa_list)
                    table.rows[batch_count].cells[2].text = str(count_less_60)
                    table.rows[batch_count].cells[3].text = str(count_greater_60)
                    table.rows[batch_count].cells[4].text = str(count_greater_90)
        doc.save(word_file_path)
        return f"{self.file_name}.docx"