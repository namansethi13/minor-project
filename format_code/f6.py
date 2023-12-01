import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor


class f1:
    def __init__(self, file_data, subjects_to_consider, classes):
        self.subjects_to_consider = subjects_to_consider
        self.file_data = file_data
        for i in self.file_data:
            # Add "S.no","Enrollment No.","Name" to the list of columns to be considered
            self.file_data[i] = ["S.no", "Enrollment No.",
                                 "Name"] + self.file_data[i]

        self.classes = classes
        self.df = pd.DataFrame()
        self.filtered_df = pd.DataFrame()
        self.all_subjects = {
            '020102': 'Applied Maths',
            '020104': 'Web Based Programming',
            '020106': 'Data Structures & Algorithm Using C',
            '020108': 'DBMS',
            '020110': 'EVS',
            '020136': 'SAUE',
            '020172': 'Practical IV-WBP Lab',
            '020174': 'Practical- V DS Lab',
            '020176': 'Practical- VI DBMS Lab',
            '020202': 'Computer Networks',
            '020204': 'Operating Systems',
            '020206': 'Computer Graphics',
            '020208': 'Software Engineering',
            '020210': 'Business Communication',
            '020236': 'SAUE',
            '020272': 'Practical- VII CN Lab',
            '020274': 'Practical- VIII OS Lab',
            '020276': 'Practical- IX CG Lab',
        }

    def process_data(self):

        for file_name, column_names in self.file_data.items():
            df1 = pd.read_excel(file_name, skiprows=6)
            df1 = df1.iloc[:-10, :-4]

            selected_columns = list(range(3)) + \
                list(range(3, len(df1.columns), 3))
            df1 = df1.iloc[:, selected_columns]
            df1.columns = column_names

            # keep only those rows which have "A" in the "A" column or "B" in the "B" column else nothing
            if 'A' in column_names:
                df1 = df1[df1['A'] == 'A']
            elif 'B' in column_names:
                df1 = df1[df1['B'] == 'B']
            else:
                continue

            self.df = pd.concat([self.df, df1], ignore_index=True)

        # keep only "Enrollment No.", "Name" and the subject in the DATAFRAME and order by subject
        self.df = self.df[['Enrollment No.', 'Name'] +
                          self.subjects_to_consider]
        self.df = self.df[(
            self.df[self.subjects_to_consider] != 0).all(axis=1)]
        self.df = self.df.sort_values(by=self.subjects_to_consider)

        # keep details of only top 10 students and bottom 10 students of each subject
        top_bottom_df = pd.DataFrame()
        for subject in self.subjects_to_consider:
            top_bottom_df = top_bottom_df._append(
                self.df.nlargest(10, subject))
            top_bottom_df = top_bottom_df._append(
                self.df.nsmallest(10, subject))

        self.df = top_bottom_df

        self.df.to_csv('f6.csv')

    def write_to_doc(self):

        doc = Document()
        for section in doc.sections:

            section.left_margin = section.right_margin = Inches(0.4)

            section.top_margin = section.bottom_margin = Inches(0)
        header_lines = [
            "",
            'MAHARAJA SURAJMAL INSTITUTE',
            'DEPARTMENT OF COMPUTER APPLICATIONS [M]',
            '                      Faculty Name: Dr. ABC       Aug-Dec 2019              Date: ',
        ]
        footer_lines = ["Faculty                         Result Analysis Committee	             HOD, BBA (M)",
                        "Ms. XYZ"]
        for line in header_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        index = 0
        for i in self.subjects_to_consider:
            # leave a few lines
            doc.add_paragraph()
            doc.add_table(rows=13, cols=7)
            table = doc.tables[-1]
            table.style = "Table Grid"
            table.autofit = True
            # set the width of each column to take up the space of the largest word in that column

            # put Subject name in the first row
            table.cell(
                0, 1).text = f'Subject Name: {self.all_subjects[i]}    Class:{self.classes[index]}'
            table.cell(0, 1).merge(table.cell(0, 6))
            # put "Top 10" in the second row's 2nd column and "Bottom 10" in the second row's 5th column
            table.cell(1, 1).text = "Top 10 Students"
            table.cell(1, 4).text = "Bottom 10 Students"
            # merge the cells in the second row
            table.cell(1, 1).merge(table.cell(1, 3))
            table.cell(1, 4).merge(table.cell(1, 6))
            # put S.no in the third row's 1st column, Enrollment No. in the third row's 2nd column, Name in the third row's 3rd column, Marks in the third row's 4th column, Enrollment No. in the third row's 5th column, Name in the third row's 6th column, Marks in the third row's 7th column
            table.cell(2, 0).text = "S.no"
            table.cell(2, 1).text = "Enrol No."
            table.cell(2, 2).text = "Name"
            table.cell(2, 3).text = "Marks"
            table.cell(2, 4).text = "Enrol No."
            table.cell(2, 5).text = "Name"
            table.cell(2, 6).text = "Marks"
            # for each 20 entries in the dataframe put the details of first 10 students in 2-4 columns and the details of next 10 students in 5-7 columns

            for j in range(10):
                table.cell(3+j, 0).text = str(j+1)
                table.cell(
                    3+j, 1).text = f"{int(self.df.iloc[j+20*index, 0]):011d}"
                table.cell(
                    3+j, 2).text = str(self.df.iloc[j+20*index, 1]).casefold().title()
                table.cell(
                    3+j, 3).text = str(int(self.df.iloc[j+20*index, 2+index]))
                table.cell(
                    3+j, 4).text = f"{int(self.df.iloc[j+10+20*index, 0]):011d}"
                table.cell(
                    3+j, 5).text = str(self.df.iloc[j+10 + 20*index, 1]).casefold().title()
                table.cell(
                    3+j, 6).text = str(int(self.df.iloc[j+10 + 20*index, 2+index]))

            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.space_before = Pt(0)
                        paragraph_format.space_after = Pt(0)
                        paragraph_format.line_spacing = 1
                        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            font = run.font
                            font.name = 'Times New Roman'
                            font.size = Pt(11)
                            font.bold = True
                            font.color.rgb = RGBColor(0, 0, 0)
            index += 1
        doc.add_paragraph()
        for line in footer_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


# Set the font to Times New Roman
        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(
                    20) if paragraph.text == 'MAHARAJA SURAJMAL INSTITUTE' else Pt(14)
                font.bold = True
                if paragraph.text == 'Class-wise Result Analysis':
                    font.underline = True
                font.color.rgb = RGBColor(0, 0, 0)
        doc.save("f6.docx")


file_data = {
    "sem2.xlsx": ["B", '020102', '020104', '020106', '020108', '020110', '020136', '020172', '020174', '020176'],
    "sem3.xlsx": ["A", '020202', '020204', '020206', '020208', '020210', '020236', '020272', '020274']
}
needed_subjects = ['020102', '020204']
classes = ["BCA 2B (M)", "BCA 3A (M)"]
f1 = f1(file_data, needed_subjects, classes)
f1.process_data()
f1.write_to_doc()