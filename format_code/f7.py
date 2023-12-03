import docx
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.shared import Pt, RGBColor
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT


class f7:
    def __init__(self,input_file, file_data):
        self.file_data = file_data
        self.input_file = input_file
        self.file_data["Subjects"] = ["SNo","Enrollment No.", "Name","Section"] + self.file_data["Subjects"]+["CGPA%"]
        self.df = pd.DataFrame()
        self.filtered_df = pd.DataFrame()
        self.all_subjects = {
            '020102': 'Applied Maths',
            '020104': 'Web Based Programming',
            '020106': 'Data Structures & Algorithm Using C',
            '020108': 'DBMS',
            '020110': 'EVS',
            '020136': 'SAUE',
            '020172': 'Practical IV-WBP Lab',
            '020174': 'Practical- V DS Lab',
            '020176': 'Practical- VI DBMS Lab',
            '020202': 'Computer Networks',
            '020204': 'Operating Systems',
            '020206': 'Computer Graphics',
            '020208': 'Software Engineering',
            '020210': 'Business Communication',
            '020236': 'SAUE',
            '020272': 'Practical- VII CN Lab',
            '020274': 'Practical- VIII OS Lab',
        }
    def write_to_doc(self):
        self.df = pd.read_excel(self.input_file,skiprows=5)
        #remove last 10 rows and last 2 columns and 4th last column
        #sort by cgpa

        self.df = self.df.iloc[:-10,:-2]
        self.df = self.df.drop(self.df.columns[-2],axis=1)
        #rename columns 1-4 as SNo, Enrollment No., Name, Section
        self.df.rename(columns={'Unnamed: 0':'SNo',' ':'Enrollment No.','Unnamed: 2':'Name','Unnamed: 3':'Section'},inplace=True)
        self.df = self.df.sort_values(by=['CGPA%'],ascending=False)
        #only keep those rows which have section A
        self.df = self.df[self.df['Section']==self.file_data['Section']]
        #remove section column
        self.df = self.df.drop('Section',axis=1)
        #remove NaN values
        self.df = self.df.dropna()
        #Now make a new dataframe with top 10 students
        self.topdf = self.df.iloc[:10,:]
        #Now make a new dataframe with bottom 10 students
        self.bottomdf = self.df.iloc[-10:,:]
        self.bottomdf=self.bottomdf.sort_values(by=['CGPA%'],ascending=True)
        print(self.topdf)
        print(self.bottomdf)
        doc=Document()
        def change_orientation():
            current_section = doc.sections[-1]
            new_width, new_height = current_section.page_height, current_section.page_width
            new_section = doc.add_section(WD_SECTION.NEW_PAGE)
            new_section.orientation = WD_ORIENT.LANDSCAPE
            new_section.page_width = new_width
            new_section.page_height = new_height
            return new_section

        change_orientation()
        header_lines = [
            
            'MAHARAJA SURAJMAL INSTITUTE',
            'DEPARTMENT OF COMPUTER APPLICATIONS [M]',
            '                      Faculty Name: Dr. ABC       Aug-Dec 2019              Date: ',
        ]
        footer_lines = ["Faculty                         Result Analysis Committee	             HOD, BBA (M)",
                        "Ms. XYZ"]
        
        for section in doc.sections:
            section.left_margin = section.right_margin = Inches(0.4)
            section.top_margin = section.bottom_margin = Inches(0)
            section.orientation = WD_ORIENTATION.LANDSCAPE
        
        for line in header_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_table(rows=13,cols=(self.file_data["Subjects"].__len__()-5)*3+4)
        table = doc.tables[-1]
        table.style = 'Table Grid'
        table.cell(0,0).text = "SNo"
        table.cell(0,1).text = "Enrollment No."
        table.cell(0,2).text = "Name of the student"
        for i in range(len(self.file_data["Faculty Names"])):

            table.cell(0,3+i*3).text = self.file_data["Faculty Names"][i]
            table.cell(0,4+i*3).text = ""
            table.cell(0,5+i*3).text = ""
            table.cell(0,3+i*3).merge(table.cell(0,5+i*3))
        for i in range(3):
            table.cell(1,i).text = ""
        for i in range(len(self.file_data["Faculty Names"])):
            table.cell(1,3+i*3).text = self.all_subjects[self.file_data["Subjects"][4+i]]
            table.cell(1,4+i*3).text = ""
            table.cell(1,5+i*3).text = ""
            table.cell(1,3+i*3).merge(table.cell(1,5+i*3))
        table.cell(1,3+len(self.file_data["Faculty Names"])*3).text = "Total"
        table.cell(1,3+len(self.file_data["Faculty Names"])*3).merge(table.cell(0,3+len(self.file_data["Faculty Names"])*3))
        table.cell(2,3+len(self.file_data["Faculty Names"])*3).text = "%"
        for i in range(3):
            table.cell(2,i).text = ""
        for i in range(len(self.file_data["Faculty Names"])):
            if self.file_data["Subjects"][4+i] in self.file_data["Practicals"]:
                table.cell(2,3+i*3).text = "Int (40)"
                table.cell(2,4+i*3).text = "Ext (60)"
                table.cell(2,5+i*3).text = "T (100)"
            else:
                table.cell(2,3+i*3).text = "Int (25)"
                table.cell(2,4+i*3).text = "Ext (75)"
                table.cell(2,5+i*3).text = "T (100)"
        #print the dataframe in the table
        for i in range(10):
            table.cell(i+3,0).text = str(i+1)
            table.cell(i+3,1).text = f"{self.topdf.iloc[i,1]:.0f}"
            table.cell(i+3,2).text = str(self.topdf.iloc[i,2]).title()
            for j in range(len(self.file_data["Faculty Names"])):
                table.cell(i+3,3+j*3).text = f"{self.topdf.iloc[i,3+j*3]:.0f}"
                table.cell(i+3,4+j*3).text = f"{self.topdf.iloc[i,4+j*3]:.0f}"if self.topdf.iloc[i,4+j*3]!="0" else "A"
                table.cell(i+3,5+j*3).text = f"{self.topdf.iloc[i,5+j*3]:.0f}"if self.topdf.iloc[i,4+j*3]!="0" else "A"
            table.cell(i+3,3+len(self.file_data["Faculty Names"])*3).text = f"{self.topdf.iloc[i,-1]:.2f}%"
            for row in table.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.space_before = Pt(0)
                            paragraph_format.space_after = Pt(0)
                            paragraph_format.line_spacing = 1
                            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                font = run.font
                                font.name = 'Times New Roman'
                                font.size = Pt(11)
                                font.bold = True
                                font.color.rgb = RGBColor(0, 0, 0)
        
        for line in footer_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_page_break()
        
        for line in header_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_table(rows=13,cols=(self.file_data["Subjects"].__len__()-5)*3+4)
        table = doc.tables[-1]
        table.style = 'Table Grid'
        table.cell(0,0).text = "SNo"
        table.cell(0,1).text = "Enrollment No."
        table.cell(0,2).text = "Name of the student"
        for i in range(len(self.file_data["Faculty Names"])):

            table.cell(0,3+i*3).text = self.file_data["Faculty Names"][i]
            table.cell(0,4+i*3).text = ""
            table.cell(0,5+i*3).text = ""
            table.cell(0,3+i*3).merge(table.cell(0,5+i*3))
        for i in range(3):
            table.cell(1,i).text = ""
        for i in range(len(self.file_data["Faculty Names"])):
            table.cell(1,3+i*3).text = self.all_subjects[self.file_data["Subjects"][4+i]]
            table.cell(1,4+i*3).text = ""
            table.cell(1,5+i*3).text = ""
            table.cell(1,3+i*3).merge(table.cell(1,5+i*3))
        table.cell(1,3+len(self.file_data["Faculty Names"])*3).text = "Total"
        table.cell(1,3+len(self.file_data["Faculty Names"])*3).merge(table.cell(0,3+len(self.file_data["Faculty Names"])*3))
        table.cell(2,3+len(self.file_data["Faculty Names"])*3).text = "%"
        for i in range(3):
            table.cell(2,i).text = ""
        for i in range(len(self.file_data["Faculty Names"])):
            if self.file_data["Subjects"][4+i] in self.file_data["Practicals"]:
                table.cell(2,3+i*3).text = "Int (40)"
                table.cell(2,4+i*3).text = "Ext (60)"
                table.cell(2,5+i*3).text = "T (100)"
            else:
                table.cell(2,3+i*3).text = "Int (25)"
                table.cell(2,4+i*3).text = "Ext (75)"
                table.cell(2,5+i*3).text = "T (100)"
        #print the dataframe in the table
        for i in range(10):
            table.cell(i+3,0).text = str(i+1)
            table.cell(i+3,1).text = f"{self.bottomdf.iloc[i,1]:.0f}"
            table.cell(i+3,2).text = str(self.bottomdf.iloc[i,2]).title()
            for j in range(len(self.file_data["Faculty Names"])):
                table.cell(i+3,3+j*3).text = f"{self.bottomdf.iloc[i,3+j*3]:.0f}"
                table.cell(i+3,4+j*3).text = f"{self.bottomdf.iloc[i,4+j*3]:.0f}"if self.bottomdf.iloc[i,4+j*3]!=0 else "A"
                table.cell(i+3,5+j*3).text = f"{self.bottomdf.iloc[i,5+j*3]:.0f}"if self.bottomdf.iloc[i,5+j*3]!=0 else "A"
            table.cell(i+3,3+len(self.file_data["Faculty Names"])*3).text = f"{self.bottomdf.iloc[i,-1]:.2f}%"
        for line in footer_lines:
            paragraph = doc.add_heading(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #style the table
        for row in table.rows:
                    for cell in row.cells:
                        paragraphs = cell.paragraphs
                        for paragraph in paragraphs:
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.space_before = Pt(0)
                            paragraph_format.space_after = Pt(0)
                            paragraph_format.line_spacing = 1
                            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for run in paragraph.runs:
                                font = run.font
                                font.name = 'Times New Roman'
                                font.size = Pt(9)
                                font.bold = True
                                font.color.rgb = RGBColor(0, 0, 0)
        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(
                    20) if paragraph.text == 'MAHARAJA SURAJMAL INSTITUTE' else Pt(14)
                font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)
        doc.save("test.docx")
        
file_data = {
    
        "Subjects": ["020102", "020104", "020106", "020108", "020110", "020136", "020172", "020174", "020176"],
        "Faculty Names": ["Dr. ABC", "Dr. DEF", "Dr. GHI", "Dr. JKL", "Dr. MNO", "Dr. PQR", "Dr. STU", "Dr. VWX", "Dr. YZ"],
        "Practicals": ["020172", "020174", "020176"],
        "Section":'A',
        "Semester": '2',
    
}
f7= f7("sem2.xlsx",file_data)
f7.write_to_doc()