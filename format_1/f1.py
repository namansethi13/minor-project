
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt,RGBColor


class Format_1:

    def __init__(self, input_file,all_subjects):
        self.input_file = input_file
        self.all_subjects = all_subjects
        self.df = None
        self.excel_df = None
        self.filtered_df = pd.DataFrame(columns=['Paper Code', 'Subjects Taught', 'Students Appeared', 'Passed',
                                        'Pass %', '>=90%', '89.99 - 75%', '74.99 - 60%', '59.99 - 50%', '49.99-40%', '<40%', 'Highest Marks'])
        self.excel_file_path = None
        self.word_file_path = None
        self.subject_name_mapping = {'S.No': 'S.No',
                                     'Enrollment Number': 'Enrollment Number',
                                     'Name': 'Name',
                                     'Sec': 'Sec',
                                     'Total Marks': 'Total Marks',
                                     'CGPA%': 'CGPA%',
                                     'Reappear': 'Reappear',
                                     'Absent': 'Absent',}
        for key, value in self.all_subjects.items():
            self.subject_name_mapping.update({key: key+" "+value+" (Internal)"})
            self.subject_name_mapping.update({key + ".1": key+" "+value + " (External)"})
            self.subject_name_mapping.update({key + ".2": key+" "+value + " (Total)"})

    def read_data(self,subjects,students,):

        col_names = ['S.No', 'Enrollment Number', 'Name', 'Sec',]
        for i in range(len(subjects)):
            col_names += [subjects[i], subjects[i]+".1", subjects[i]+".2"]
        col_names += [
                     'Total Marks', 'CGPA%',  'Reappear', 'Absent',]
        self.df = pd.read_excel(
            self.input_file, skiprows=6, nrows=students+1, names=col_names)
        self.df.columns = [self.subject_name_mapping.get(
            col) for col in self.df.columns]

    def filter_data(self, subjects, sec):
        # Start with the fixed columns
        columns = ['Sec']

    # Add the subject columns
        for subject in subjects:
            columns.append(self.subject_name_mapping[subject + '.2'])

    # Filter the DataFrame
        self.excel_df = self.df[columns]
        self.excel_df = self.excel_df[self.excel_df['Sec'] == sec]
        self.excel_df = self.excel_df.drop(columns=['Sec'])

    def write_to_excel(self):
        self.excel_file_path = 'output2.xlsx'
        self.excel_df.to_excel(self.excel_file_path, index=False)
        print("Excel file created at: ", self.excel_file_path)

    def read_from_filtered_excel(self, course_name,subject_code):
        
        self.filtered_df = pd.DataFrame()  # Initialize filtered_df as an empty dataframe

    # Assuming the column you want to work with is named 'YourColumnName'
        column_names = list()
        for i in range(len(subject_code)):
            column_names.append(subject_code[i]+" "+self.all_subjects[subject_code[i]]+" (Total)")
        
                        
        
        for i, column_name in enumerate(column_names):
            self.filtered_df.at[i, 'S.No'] = f'{i+1:.0f}'
            self.filtered_df.at[i, 'Paper Code'] = str(
                course_name + str(column_name[3:6]))
            self.filtered_df.at[i, 'Subjects Taught'] = str(
                column_name[7:-8])

            total_students = self.excel_df[self.excel_df[column_name]
                                           != 0].shape[0]
            self.filtered_df.at[i, 'Students Appeared'] = f'{(total_students):.0f}'

            passed_students = self.excel_df[self.excel_df[column_name]
                                            >= 40].shape[0]
            self.filtered_df.at[i, 'Passed'] = f'{(passed_students):.0f}'
            self.filtered_df.at[i,
                                'Pass %'] = f"{(passed_students / total_students) * 100:.2f}"
            countA1 = self.excel_df[self.excel_df[column_name] >= 90].shape[0]
            countA2 = self.excel_df[(self.excel_df[column_name] >= 75) & (
                self.excel_df[column_name] <= 89)].shape[0]
            countA3 = self.excel_df[(self.excel_df[column_name] >= 60) & (
                self.excel_df[column_name] <= 74)].shape[0]
            countB1 = self.excel_df[(self.excel_df[column_name] >= 50) & (
                self.excel_df[column_name] <= 59)].shape[0]
            countB2 = self.excel_df[(self.excel_df[column_name] >= 40) & (
                self.excel_df[column_name] <= 49)].shape[0]
            countB3 = self.excel_df[(self.excel_df[column_name] >= 1) & (
                self.excel_df[column_name] <= 39)].shape[0]
            self.filtered_df.at[i, '>=90%'] = str(
                f"{countA1}\n({countA1 / total_students * 100:.2f})")
            self.filtered_df.at[i,
                                '89.99 - 75%'] = f"{countA2}\n({countA2 / total_students * 100:.2f})"
            self.filtered_df.at[i,
                                '74.99 - 60%'] = f"{countA3}\n({countA3 / total_students * 100:.2f})"
            self.filtered_df.at[i,
                                '59.99 - 50%'] = f"{countB1}\n({countB1 / total_students * 100:.2f})"
            self.filtered_df.at[i,
                                '49.99-40%'] = f"{countB2}\n({countB2 / total_students * 100:.2f})"
            self.filtered_df.at[i,
                                '<40%'] = f"{countB3}\n({countB3 / total_students * 100:.2f})"
            self.filtered_df.at[i,'C=B-A'] = f'{(countB1+countB2+countB3-countA1):.0f}'
            self.filtered_df.at[i,
                                'Highest Marks'] = f'{self.excel_df[column_name].max():.0f}'


    def write_to_doc(self):
        self.word_file_path = 'output.docx'
        excel_data_df = self.filtered_df
        doc = Document()
        for section in doc.sections:
            
            section.left_margin =section.right_margin= Inches(0.4)
            
            section.top_margin = section.bottom_margin = Inches(0)
            
# Add lines to the header
        header_lines = [
            "",
            'Maharaja Surajmal Institute',
            'Department of Computer Applications [M]',
            'Date: …………',
            'Faculty Name: - Dr. ABC                        Shift-I                                Max Marks: 100 ',
            'Result Analysis (Aug-Dec 2019)'
        ]

        for line in header_lines:
            paragraph = doc.add_heading(line)
            if line == 'Date: …………':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif line == 'Result Analysis (Aug-Dec 2019)':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif line == 'Faculty Name: - Dr. ABC                        Shift-I                                Max Marks: 100 ':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# Set the font to Times New Roman
        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(20) if paragraph.text == 'Maharaja Surajmal Institute' else Pt(14)
                font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)  # RGB values for black
        
            
        doc.add_table(
            excel_data_df.shape[0]+1, excel_data_df.shape[1], style='Table Grid')
        table = doc.tables[0]
        # add the header rows.
        for j in range(excel_data_df.shape[-1]):
            cell = table.cell(0, j)
            cell.text = excel_data_df.columns[j]
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)

# Add the rest of the data frame with bold, Times New Roman, and text size 10
        for i in range(excel_data_df.shape[0]):
            for j in range(excel_data_df.shape[-1]):
                cell = table.cell(i + 1, j)
                cell.text = str(excel_data_df.values[i, j])
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
                cell.paragraphs[0].runs[0].font.size = Pt(10)
        #add a row at the end 		Total Students & Pass %	118	117	99.15%	No. of students & average % above 60% 82 (69.49%)	No. of students & average % below 60% 34 (28.81%)	

        # table.add_row()
        # cell = table.cell(excel_data_df.shape[0]+1, 0)
        # cell.text = "Total Students & Pass %"
        # cell.paragraphs[0].runs[0].bold = True
        # cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        # cell.paragraphs[0].runs[0].font.size = Pt(10)
        # cell = table.cell(excel_data_df.shape[0]+1, 1)
        # appeared_text = f'{excel_data_df["Students Appeared"].sum():.0f}'
        # cell.text = appeared_text
        # cell.paragraphs[0].runs[0].bold = True
        # cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        # cell.paragraphs[0].runs[0].font.size = Pt(10)
        # cell = table.cell(excel_data_df.shape[0]+1, 2)
        # cell.text = f'{excel_data_df["Passed"].sum():.0f}'
        # cell.paragraphs[0].runs[0].bold = True
        # cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        # cell.paragraphs[0].runs[0].font.size = Pt(10)
        # cell = table.cell(excel_data_df.shape[0]+1, 3)
        # cell.text = f'{(excel_data_df["Passed"].sum()/excel_data_df["Students Appeared"].sum())*100:.2f}%'
        # cell.paragraphs[0].runs[0].bold = True
        # cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        # cell.paragraphs[0].runs[0].font.size = Pt(10)
        # cell = table.cell(excel_data_df.shape[0]+1, 4)
        # cell.text = "No. of students & average % above 60%"
        # cell.paragraphs[0].runs[0].bold = True
        # cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        # cell.paragraphs[0].runs[0].font.size = Pt(10)
        # cell = table.cell(excel_data_df.shape[0]+1, 5)
        # cell.text = f'{excel_data_df["Students Appeared"].sum()-excel_data_df[excel_data_df["<40%"]!="0"].shape[0]:.0f} ({(excel_data_df["Students Appeared"].sum()-excel_data_df[excel_data_df["<40%"]!="0"].shape[0])/excel_data_df["Students Appeared"].sum()*100:.2f}%)'
        # cell.paragraphs[0].runs[0].bold = True
        # cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        # cell.paragraphs[0].runs[0].font.size = Pt(10)
        # cell = table.cell(excel_data_df.shape[0]+1, 6)
        # cell.text = "No. of students & average % below 60%"
        # cell.paragraphs[0].runs[0].bold = True
        # cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        # cell.paragraphs[0].runs[0].font.size = Pt(10)
        # cell = table.cell(excel_data_df.shape[0]+1, 7)
        # cell.text = f'{excel_data_df[excel_data_df["<40%"]!="0"].shape[0]:.0f} ({excel_data_df[excel_data_df["<40%"]!="0"].shape[0]/excel_data_df["Students Appeared"].sum()*100:.2f}%)'
        # cell.paragraphs[0].runs[0].bold = True
        # cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        # cell.paragraphs[0].runs[0].font.size = Pt(10)
        
# Add the footer
        footer_lines = [
            "",
            f'C= No. of Students securing 90 or above is deducted from the total no. of students securing below 60 marks and accordingly the %age below 60% (aggregate) is computed ',
            '',
            '',
            '',
            '“I do hereby solemnly affirm and declare that the facts stated in the above result are true to the best of my knowledge and belief”',
            """Dr.ABC     		                 (Dr.Menal Dahiya)       				(Mr. Manoj Kumar)	
Assistant Professor 		Convenor-Result Analysis Committee         	HOD-BCA[M]"""
        ]
        for line in footer_lines:
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(10)
                if line != '“I do hereby solemnly affirm and declare that the facts stated in the above result are true to the best of my knowledge and belief”':
                    font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)  # RGB values for black
            

        
        doc.save(self.word_file_path)


#driver code
subjects=['020102', '020104', '020106','020108', '020110', '020136','020172','020174','020176']
all_subjects={'020102': 'Applied Maths',
              '020104': 'Web Based Programming',
              '020106': 'Data Structures & Algorithm Using C',
              '020108': 'DBMS',
              '020110': 'EVS',
              '020136': 'SAUE',
              '020172': 'Practical IV-WBP Lab',
              '020174': 'Practical- V DS Lab',
              '020176': 'Practical- VI DBMS Lab'}
f1 = Format_1("output.xlsx",all_subjects)
needed_subjects = ['020102', '020104', ]
f1.read_data(subjects, 111)
f1.filter_data(subjects, "B")
f1.read_from_filtered_excel("BCA",needed_subjects)
f1.write_to_doc()
